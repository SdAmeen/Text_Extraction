from flask import Flask, request, send_file, jsonify, render_template
import os
from PIL import Image
import pytesseract
from docx import Document
from googletrans import Translator

app = Flask(__name__)

# Path to the tesseract executable (modify if needed)
pytesseract.pytesseract.tesseract_cmd = r'/usr/bin/tesseract'

def extract_text_from_image(image_path):
    """Extract text from an image using Tesseract OCR."""
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img)
        return text
    except Exception as e:
        return f"Error processing image: {str(e)}"

def translate_text(text, languages):
    """Translate text into the specified languages."""
    translator = Translator()
    translations = {}
    for lang in languages:
        try:
            translated = translator.translate(text, dest=lang)
            translations[lang] = translated.text
        except Exception as e:
            translations[lang] = f"Error translating to {lang}: {str(e)}"
    return translations

def save_translations_to_word(translations, output_path):
    """Save translations to a Word document, each language on a new page."""
    doc = Document()
    for lang, translated_text in translations.items():
        doc.add_heading(f'Translation in {lang}', level=1)
        doc.add_paragraph(translated_text)
        doc.add_page_break()
    
    doc.save(output_path)

@app.route('/')
def index():
    """Render the HTML form."""
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_image():
    """Handle image upload and processing."""
    file = request.files.get('file')
    languages = request.form.getlist('language')  # Handle multiple languages
    
    if file and languages:
        # Check if the file is an image
        if file.filename.lower().endswith(('.png', '.jpg', '.jpeg')):
            image_path = os.path.join('uploads', file.filename)
            file.save(image_path)
            
            # Extract text and perform translation
            text = extract_text_from_image(image_path)
            translations = translate_text(text, languages)
            
            # Save translations to Word document
            output_path = 'translated_output.docx'
            save_translations_to_word(translations, output_path)
            
            # Send the Word document back to the user
            return send_file(output_path, as_attachment=True)
        else:
            return jsonify({'error': 'Unsupported file type. Please upload an image file.'}), 400
    
    return jsonify({'error': 'Invalid file or language'}), 400


if __name__ == '__main__':
    if not os.path.exists('uploads'):
        os.makedirs('uploads')
    app.run(debug=True)

